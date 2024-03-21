from django.shortcuts import render, HttpResponse,redirect
import pandas as pd
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .forms import UploadWorkSheetForm, UserPermissionForm, TemplateForm,FormSheet,RelationTempForm
from .models import UploadTemplate,UserPermission
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth import login as auth_login, authenticate
from django.contrib import messages
from docx import Document
import docx
import pdfplumber
import sys
from .forms import UserPermissionForm, UploadWorkSheetForm
import os
from .models import UploadTemplate, UploadedForm,TempFormRelation
from django.conf import settings
from pathlib import Path
from django.http import JsonResponse
import json
# Create your views here.
from django.shortcuts import render

from .models import FormData,TempleData
def loginuser(request):
    if request.method == 'POST':
        form = AuthenticationForm(request.POST)
        if form.is_valid:
            username = request.POST['username']
            password = request.POST['password']
            user = authenticate(username=username, password=password)
            if user:
                auth_login(request, user)
                if user.is_superuser:
                    return redirect('home')  # Superuser, redirect to the 'home' page
                else:
                    return redirect('home')  
                
            else:
                if username == '' or password == '':
                    messages.error(
                        request, message='Please Enter Username and Passowrd Correctly')
                else:
                    messages.error(
                        request, message='Username or Password not correct')

    else:
        form = AuthenticationForm()
    return render(request, 'login.html', {'form': form})


def home(request):
    return render(request,"home.html")



def addtemplate(request):
    form = UploadWorkSheetForm()

    if request.method == 'POST':
        form = UploadWorkSheetForm(request.POST, request.FILES)
        if form.is_valid():
            
            upload_template_instance = form.save()
            template_master = upload_template_instance.worksheet_file.path
            print(template_master)
            if template_master.endswith('.docx'):
                t_id = upload_template_instance.id
                print(t_id)
                # Parse the Word document using python-docx
                doc = Document(template_master)
                header = doc.sections[0].header
                headers_data = []
                for table in header.tables:
                    header_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        unique_values = list(set(row_data))
                        # print(unique_values,'header row')
                        template_form_instance = TempleData(temp_data=unique_values,template_master_id=t_id)
                        template_form_instance.save()
                        header_data.append(unique_values)
                    headers_data.append(header_data)
                tables_data = []
                columns = None
                # Extract data from each table
                for table in doc.tables:
                    table_data = []
                    unique_values = set()  # Track unique values across top 3 rows
                    # Counter for appending "row" labels
                    row_counter = 1
                    for i, row in enumerate(table.rows):
                        row_data = []
                        for j, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()  # Strip whitespace from cell text
                            # Check if it's the header row
                            if i < 3:
                                # Check for duplicate values in top 3 rows
                                if cell_text in unique_values:
                                    cell_text = ""  # Fill null value if duplicate
                                else:
                                    unique_values.add(cell_text)
                            row_data.append(cell_text)
                        table_data.append(row_data)
                        if i == 2 and '20 µ' in row_data and 'G4' in row_data:
                            row_data[-1] = 'M6(F6)/F9'
                            row_data[-2] = 'G4'
                            row_data[-3] = '20 µ'
                        if i == 3 and '20 µ' in row_data or 'G4' in row_data:
                            row_data[0]=""
                            row_data[1]=""
                        # Print the row label
                        # print(f"row{row_counter}", row_data)
                        template_form_instance = TempleData(temp_data=row_data,template_master_id=t_id)
                        template_form_instance.save()
                        row_counter += 1
                    tables_data.append(table_data)
            return redirect('selecttemplate')

        else:
            # Handle form errors - you can print them or log them for debugging
            print(form.errors)


    name=UploadTemplate.objects.all()

    context = {
        'form': form
    }
    return render(request, 'addtemplate.html', context)

def selecttemplate(request):
    tform=UserPermissionForm
    if request.method == 'POST':
        count_form = UserPermissionForm(request.POST, request.FILES)
        if count_form.is_valid():

            count_form.save()
            return redirect('formuploaded')
        else:
            print(count_form.errors)
    else:
        count_form = UserPermissionForm()
    context = {
        'tform': tform}
    return render(request, 'selecttemp.html', context)

def formuploaded(request):
    tform=FormSheet
    if request.method == 'POST':
        count_form = FormSheet(request.POST, request.FILES)
        if count_form.is_valid():   
            count_form.save()
            return redirect('formdata')
        else:
            print(count_form.errors)
    else:
        count_form = FormSheet()
    context = {
        'tform': tform} 
    return render(request, 'addform.html', context)

def editform(request):
    fid = UploadedForm.objects.latest('id')
    tid = UploadTemplate.objects.latest('id')
    tempid = tid.id
    username = tid.creator

    formdata = TempleData.objects.filter(template_master_id=tempid)

    rows = []
    for data in formdata:
        rows.append(eval(data.temp_data))

    df = pd.DataFrame(rows)

    headers = df.iloc[0]
    # df = df[1:]
    df.columns = headers

    for index, row in df.iterrows():
        # Create a new FormData object for each row
        FormData.objects.create(
            form_data=row.tolist(),
            creator=username,
            form_name=fid,
            template_master=tid,
        )

    return render(request, 'formdata.html', {'df': df})

def update(request):

    if request.method == 'POST':
        level_codes = request.POST.getlist('std_value[]')
        print(level_codes,'level codes')
        index_list = request.POST.getlist('id[]')
        print(index_list)
        for index, std_value in zip(index_list, level_codes):
            try:
                stddata = FormData.objects.get(id=int(index))
                stddata.std_value = std_value
                stddata.save()
                print(stddata,'stddaa')
            except FormData.DoesNotExist:
                # Handle the case where the LevelCode object with the given 'id' doesn't exist
                pass
    context = {
        'update_lc': update}

    return render(request,'updateform.html',context) 
def selectform(request):
    tform=RelationTempForm
    if request.method == 'POST':
        count_form = RelationTempForm(request.POST, request.FILES)
        if count_form.is_valid():   
            count_form.save()
            return redirect('showform')
        else:
            print(count_form.errors)
    else:
        count_form = RelationTempForm()
    context = {
        'tform': tform} 
    return render(request, 'selectform.html', context)


def showform(request):
    tfr = TempFormRelation.objects.latest('id')
    tid = UploadTemplate.objects.latest('id')
    fid = UploadedForm.objects.latest('id')

    tempid=tfr.upload_template_id
    formid=tfr.form_details_id
    user=tfr.user_table_id
    print(tempid ,formid, user)

    # edittemplate=FormDetail.objects.all()
    edittemplate = FormData.objects.filter(template_master_id=tid,form_name_id=fid)
    rows = []
    for data in edittemplate:
        rows.append(eval(data.form_data))

    df = pd.DataFrame(rows)

    headers = df.iloc[0]
    df = df[1:]
    df.columns = headers

    return render(request,'showform.html',{'df': df})

def user_check(request):
    edittemplate=FormData.objects.all()
    if request.method=='POST':
        return redirect('home')
    return render(request, 'userapproved.html', {'edittemplate': edittemplate})